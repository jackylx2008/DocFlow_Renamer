from __future__ import annotations

import argparse
import base64
import hashlib
import json
import logging
import os
import re
import subprocess
import sys
import time
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import fitz
import yaml
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WORD_NS}
CHECKED_SYMBOL = "0052"
TARGET_NAME_RE = re.compile(
    r"^\d{4}-\d{2}-\d{2}_.+_质保作业申请单\.docx$", re.IGNORECASE
)
DATE_RANGE_RE = re.compile(
    r"(\d{4})\s*(?:年|[./．])\s*(\d{1,2})\s*(?:月|[./．])\s*(\d{1,2})\s*(?:日)?"
    r"\s*[~～\-—–－至到]+\s*"
    r"(?:(\d{4})\s*(?:年|[./．])\s*)?(\d{1,2})\s*(?:月|[./．])\s*(\d{1,2})\s*(?:日)?"
)
SINGLE_DATE_RE = re.compile(
    r"(\d{4})\s*(?:年|[./．])\s*(\d{1,2})\s*(?:月|[./．])\s*(\d{1,2})\s*(?:日)?"
)
INVALID_FILENAME_CHARS_RE = re.compile(r'[\\/:*?"<>|]+')
WHITESPACE_RE = re.compile(r"\s+")
WORK_TYPES = ["动火作业", "有限空间作业", "5米以上高处作业", "危大工程", "配电室接电"]
JPG_SUFFIXES = {".jpg", ".jpeg"}
AI_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png"}
WORKER_LIST_IMAGE_SUFFIX = "_工人名单"
DATE_PREFIX_RE = re.compile(r"^\d{4}-\d{2}-\d{2}")
WORD_CONTENT_NAME_RE = re.compile(r"^(\d{4}-\d{2}-\d{2})_([^_]+)_.+\.docx$", re.IGNORECASE)
PDF_SUFFIX = ".pdf"
PDF_TWELVE_DIGIT_RE = re.compile(r"\d{12}")
PDF_APPLICATION_NO_RE = re.compile(r"申请\s*编号\s*[:：]\s*([0-9０-９]{12})")
PDF_TARGET_NAME_RE = re.compile(
    r"^工程类-主体质保施工_编号：[0-9０-９]{12}\.pdf$", re.IGNORECASE
)
PDF_TARGET_NAME_PREFIX = "工程类-主体质保施工_编号："
MIN_PLAIN_PDF_CJK_CHARS = 10
OCR_PAGE_LIMIT = 2
PDF_MATCH_SEPARATOR = "；"
PDF_MATCH_CACHE_VERSION = "2"
PDF_TEXT_CACHE_NAME = ".docflow_pdf_text_cache.json"
PDF_TEXT_CACHE_VERSION = "1"
PDF_FINGERPRINT_CHUNK_SIZE = 1024 * 1024
SUMMARY_EXCEL_NAME = "质保作业申请汇总.xlsx"
MANUAL_MATCH_ENV_NAME = "manual_matches.env"
IMAGE_TEXT_PROMPT = "请识别这张图片里的所有可见中文文字。只输出识别到的文字，不要解释，不要总结。"
PDF_PAGE_TEXT_PROMPT = (
    "请识别这页 PDF 截图里的所有可见中文文字。"
    "尽量保留原文中的日期、编号、施工内容、区域和单位名称，"
    "特别要准确保留施工开始时间、施工结束时间及其字段标签。"
    "只输出识别到的文字，不要解释，不要总结。"
)
LOGGER = logging.getLogger(__name__)


class TeeStream:
    def __init__(self, *streams: Any) -> None:
        self.streams = streams
        self.encoding = getattr(streams[0], "encoding", "utf-8")
        self.errors = getattr(streams[0], "errors", "replace")

    def write(self, data: str) -> int:
        for stream in self.streams:
            stream.write(data)
        return len(data)

    def flush(self) -> None:
        for stream in self.streams:
            stream.flush()

    def isatty(self) -> bool:
        return any(getattr(stream, "isatty", lambda: False)() for stream in self.streams)


@dataclass
class Record:
    项目名称: str
    质保单位: str
    分包单位: str
    质保负责人: str
    质保负责人联系电话: str
    施工区域: str
    施工开始时间: str
    施工结束时间: str
    时长天: int | None
    施工内容: str
    施工负责人: str
    施工负责人联系电话: str
    影响改动消防设备设施: str
    影响堵塞应急疏散通道: str
    危险作业: str
    原文件名: str
    新文件名: str
    文件路径: str
    申请单文件链接: str
    申请单附图链接: str
    附件目录: str
    附件目录链接: str
    匹配PDF文件名: str
    匹配PDF文件链接: str
    处理状态: str


@dataclass
class LlamaCppConfig:
    base_url: str
    model: str
    autostart: bool
    server_path: str
    model_path: str
    mmproj_path: str
    extra_dll_dirs: list[str]
    n_gpu_layers: str
    ctx_size: str
    reasoning: str
    reasoning_budget: str

    @classmethod
    def from_repo(cls, repo_root: Path) -> "LlamaCppConfig":
        env_values = load_env_file(repo_root / "common.env")
        yaml_config = load_yaml_config(repo_root / "config.yaml")

        def value(key: str, fallback: str = "") -> str:
            return resolve_setting(
                env_values.get(key) or yaml_config.get(key.lower()),
                env_values,
                fallback,
            )

        return cls(
            base_url=value("LLAMACPP_BASE_URL", "http://127.0.0.1:8080/v1").rstrip("/"),
            model=value("LLAMACPP_MODEL", "local-model"),
            autostart=value("LLAMACPP_AUTOSTART", "true").strip().lower()
            in {"1", "true", "yes", "on"},
            server_path=value("LLAMACPP_SERVER_PATH"),
            model_path=value("LLAMACPP_MODEL_PATH"),
            mmproj_path=value("LLAMACPP_MMPROJ_PATH"),
            extra_dll_dirs=split_name_list(value("LLAMACPP_EXTRA_DLL_DIRS", "./vendor/cuda12")),
            n_gpu_layers=value("LLAMACPP_N_GPU_LAYERS", "999"),
            ctx_size=value("LLAMACPP_CTX_SIZE", "8192"),
            reasoning=value("LLAMACPP_REASONING", "off"),
            reasoning_budget=value("LLAMACPP_REASONING_BUDGET", "0"),
        )


class LlamaCppClient:
    def __init__(self, config: LlamaCppConfig, repo_root: Path) -> None:
        self.config = config
        self.repo_root = repo_root
        self._process: subprocess.Popen[Any] | None = None
        self._log_handles: list[Any] = []

    @property
    def server_root_url(self) -> str:
        parsed = urlparse(self.config.base_url)
        path = parsed.path.rstrip("/")
        if path.endswith("/v1"):
            path = path[:-3].rstrip("/")
        return f"{parsed.scheme}://{parsed.netloc}{path}".rstrip("/")

    @property
    def api_base_url(self) -> str:
        parsed = urlparse(self.config.base_url)
        if parsed.path.rstrip("/").endswith("/v1"):
            return self.config.base_url
        return f"{self.server_root_url}/v1"

    @property
    def chat_url(self) -> str:
        return f"{self.api_base_url}/chat/completions"

    @property
    def models_url(self) -> str:
        return f"{self.api_base_url}/models"

    @property
    def health_url(self) -> str:
        return f"{self.server_root_url}/health"

    def _get_json(self, url: str, timeout: float = 5) -> Any:
        with urlopen(url, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))

    def _post_json(self, url: str, payload: dict[str, Any], timeout: float = 180) -> Any:
        request = Request(
            url,
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))

    def is_server_available(self) -> bool:
        try:
            self._get_json(self.health_url, timeout=3)
            return True
        except (HTTPError, URLError, TimeoutError, OSError, json.JSONDecodeError):
            return False

    def ensure_server(self) -> None:
        if self.is_server_available():
            return
        if not self.config.autostart:
            raise RuntimeError("本地 AI 服务不可用，且 LLAMACPP_AUTOSTART 未开启")
        self.start_server()
        deadline = time.monotonic() + 180
        while time.monotonic() < deadline:
            if self.is_server_available():
                return
            time.sleep(2)
        raise RuntimeError("等待本地 AI 服务启动超时")

    def start_server(self) -> None:
        server_path = Path(self.config.server_path)
        model_path = Path(self.config.model_path)
        mmproj_path = Path(self.config.mmproj_path) if self.config.mmproj_path else None
        if not server_path.is_file():
            raise FileNotFoundError(f"LLAMACPP_SERVER_PATH 不存在: {server_path}")
        if not model_path.is_file():
            raise FileNotFoundError(f"LLAMACPP_MODEL_PATH 不存在: {model_path}")
        if mmproj_path and not mmproj_path.is_file():
            raise FileNotFoundError(f"LLAMACPP_MMPROJ_PATH 不存在: {mmproj_path}")

        parsed = urlparse(self.server_root_url)
        host = parsed.hostname or "127.0.0.1"
        port = str(parsed.port or 8080)
        command = [
            str(server_path),
            "-m",
            str(model_path),
            "--alias",
            self.config.model,
            "-c",
            self.config.ctx_size,
            "-ngl",
            self.config.n_gpu_layers,
            "--reasoning",
            self.config.reasoning,
            "--reasoning-budget",
            self.config.reasoning_budget,
            "--host",
            host,
            "--port",
            port,
        ]
        if mmproj_path:
            command[3:3] = ["--mmproj", str(mmproj_path)]

        log_dir = self.repo_root / "log"
        log_dir.mkdir(parents=True, exist_ok=True)
        stdout_handle = (log_dir / "llama_server.out.log").open("a", encoding="utf-8")
        stderr_handle = (log_dir / "llama_server.err.log").open("a", encoding="utf-8")
        self._log_handles.extend([stdout_handle, stderr_handle])

        env = os.environ.copy()
        dll_dirs = [server_path.parent]
        for raw_dir in self.config.extra_dll_dirs:
            dll_dir = Path(raw_dir)
            if not dll_dir.is_absolute():
                dll_dir = self.repo_root / dll_dir
            dll_dirs.append(dll_dir)
        env["PATH"] = os.pathsep.join(str(path) for path in dll_dirs) + os.pathsep + env.get("PATH", "")

        self._process = subprocess.Popen(
            command,
            cwd=str(server_path.parent),
            env=env,
            stdout=stdout_handle,
            stderr=stderr_handle,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )

    def assert_model_available(self) -> None:
        models = self._get_json(self.models_url, timeout=10)
        model_ids = [
            model_id
            for item in models.get("data", [])
            if isinstance(item, dict)
            for model_id in [item.get("id")]
            if isinstance(model_id, str) and model_id
        ]
        if self.config.model not in model_ids:
            raise RuntimeError(
                f"本地 AI 模型不可用: {self.config.model}; 当前模型: {', '.join(model_ids)}"
            )

    def extract_image_bytes_text(
        self,
        image_bytes: bytes,
        mime_type: str,
        prompt: str = IMAGE_TEXT_PROMPT,
        max_tokens: int = 512,
    ) -> str:
        image_b64 = base64.b64encode(image_bytes).decode("ascii")
        payload = {
            "model": self.config.model,
            "temperature": 0,
            "max_tokens": max_tokens,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt,
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{image_b64}",
                            },
                        },
                    ],
                }
            ],
        }
        response = self._post_json(self.chat_url, payload)
        choices = response.get("choices") or []
        if not choices:
            return ""
        content = choices[0].get("message", {}).get("content", "")
        if isinstance(content, list):
            return "\n".join(
                str(item.get("text", "")) for item in content if isinstance(item, dict)
            )
        return str(content or "")

    def extract_image_text(self, image_path: Path) -> str:
        mime_type = "image/png" if image_path.suffix.lower() == ".png" else "image/jpeg"
        return self.extract_image_bytes_text(
            image_path.read_bytes(),
            mime_type,
            IMAGE_TEXT_PROMPT,
            max_tokens=512,
        )

    def shutdown_server(self) -> None:
        if self._process is not None:
            self._process.terminate()
            try:
                self._process.wait(timeout=20)
            except subprocess.TimeoutExpired:
                self._process.kill()
                self._process.wait(timeout=10)
            self._process = None
        for handle in self._log_handles:
            try:
                handle.close()
            except OSError:
                pass
        self._log_handles.clear()


def load_env_file(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}

    values: dict[str, str] = {}
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def load_yaml_config(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    if not isinstance(data, dict):
        raise ValueError(f"配置文件格式不正确: {path}")
    return data


def resolve_setting(
    raw_value: str | None, env_values: dict[str, str], fallback: str
) -> str:
    if not raw_value:
        return fallback
    match = re.fullmatch(r"\$\{([A-Z0-9_]+):-([^}]*)\}", str(raw_value).strip())
    if not match:
        return str(raw_value)
    env_key, default_value = match.groups()
    return env_values.get(env_key, default_value)


def normalize_whitespace(value: str) -> str:
    return WHITESPACE_RE.sub(" ", value or "").strip()


def normalize_match_text(value: str) -> str:
    normalized = WHITESPACE_RE.sub("", value or "").strip().lower()
    return (
        normalized.replace("～", "~")
        .replace("—", "~")
        .replace("–", "~")
        .replace("－", "~")
        .replace("-", "~")
        .replace("（", "(")
        .replace("）", ")")
        .replace("，", ",")
        .replace("、", ",")
    )


def count_cjk_chars(value: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", value or ""))


def split_name_list(value: str) -> list[str]:
    raw_value = (value or "").strip()
    if raw_value.startswith("[") and raw_value.endswith("]"):
        parsed = yaml.safe_load(raw_value)
        if isinstance(parsed, list):
            return [normalize_whitespace(str(item)) for item in parsed if str(item).strip()]

    return [
        item.strip()
        for item in re.split(r"[,，；;]", raw_value)
        if item and item.strip()
    ]


def split_pdf_names(value: str) -> list[str]:
    return split_name_list(value)


def normalize_digits(value: str) -> str:
    return value.translate(str.maketrans("０１２３４５６７８９", "0123456789"))


def extract_pdf_application_no_from_name(pdf_name: str) -> str:
    stem = normalize_digits(Path(pdf_name).stem)
    if PDF_TWELVE_DIGIT_RE.fullmatch(stem):
        return stem
    if "编号" not in stem:
        return ""
    compact_digits = re.sub(r"\D", "", stem)
    match = PDF_TWELVE_DIGIT_RE.search(compact_digits)
    return match.group(0) if match else ""


def build_pdf_name_by_application_no(pdf_path_index: dict[str, Path]) -> dict[str, str]:
    pdf_name_by_application_no: dict[str, str] = {}
    for pdf_name in pdf_path_index:
        application_no = extract_pdf_application_no_from_name(pdf_name)
        if not application_no:
            continue
        existing_name = pdf_name_by_application_no.get(application_no)
        if not existing_name or pdf_name.startswith(PDF_TARGET_NAME_PREFIX):
            pdf_name_by_application_no[application_no] = pdf_name
    return pdf_name_by_application_no


def normalize_pdf_match_cache_names(
    pdf_match_cache: dict[str, str],
    pdf_path_index: dict[str, Path],
) -> dict[str, str]:
    pdf_name_by_application_no = build_pdf_name_by_application_no(pdf_path_index)
    if not pdf_name_by_application_no:
        return pdf_match_cache

    normalized_cache: dict[str, str] = {}
    renamed_count = 0
    for cache_key, pdf_names in pdf_match_cache.items():
        normalized_pdf_names: list[str] = []
        for pdf_name in split_pdf_names(pdf_names):
            application_no = extract_pdf_application_no_from_name(pdf_name)
            normalized_pdf_name = pdf_name_by_application_no.get(application_no, pdf_name)
            if normalized_pdf_name != pdf_name:
                renamed_count += 1
            normalized_pdf_names.append(normalized_pdf_name)
        normalized_cache[cache_key] = PDF_MATCH_SEPARATOR.join(normalized_pdf_names)

    if renamed_count:
        LOGGER.info("已规范化 %s 个已有 PDF 匹配缓存文件名", renamed_count)
    return normalized_cache


def sanitize_filename_part(value: str) -> str:
    cleaned = INVALID_FILENAME_CHARS_RE.sub("-", value.strip())
    cleaned = normalize_whitespace(cleaned)
    return cleaned.strip(". ")


def read_first_form_text(docx_path: Path) -> str:
    document = Document(str(docx_path))
    if not document.tables:
        raise ValueError("Word 文档中未找到表格")
    table = document.tables[0]
    if not table.rows or not table.rows[0].cells:
        raise ValueError("Word 表格结构不符合预期")
    return table.rows[0].cells[0].text


def read_first_form_runs(docx_path: Path) -> list[dict[str, str]]:
    with ZipFile(docx_path) as archive:
        xml_bytes = archive.read("word/document.xml")

    root = ET.fromstring(xml_bytes)
    table = root.find(".//w:tbl", NS)
    if table is None:
        raise ValueError("document.xml 中未找到表格")

    row = table.find("./w:tr", NS)
    if row is None:
        raise ValueError("document.xml 中未找到第一行")

    cell = row.find("./w:tc", NS)
    if cell is None:
        raise ValueError("document.xml 中未找到第一列")

    runs: list[dict[str, str]] = []
    for paragraph in cell.findall(".//w:p", NS):
        for run in paragraph.findall("./w:r", NS):
            text_nodes = run.findall("./w:t", NS)
            text = "".join(node.text or "" for node in text_nodes)
            if text:
                runs.append({"type": "text", "value": text})

            sym = run.find("./w:sym", NS)
            if sym is not None:
                char_value = sym.attrib.get(f"{{{WORD_NS}}}char", "").upper()
                if char_value:
                    runs.append({"type": "sym", "value": char_value})
    return runs


def capture_between(text: str, start_label: str, end_label: str) -> str:
    pattern = re.compile(
        rf"{re.escape(start_label)}\s*(.*?)\s*(?={re.escape(end_label)})",
        re.DOTALL,
    )
    match = pattern.search(text)
    return normalize_whitespace(match.group(1)) if match else ""


def parse_date_range(raw_value: str) -> tuple[str, str, int | None]:
    raw_value = normalize_whitespace(raw_value)
    match = DATE_RANGE_RE.search(raw_value)
    if match:
        start_year, start_month, start_day, end_year, end_month, end_day = (
            match.groups()
        )
        start_date = date(int(start_year), int(start_month), int(start_day))
        end_date = date(int(end_year or start_year), int(end_month), int(end_day))
        duration = (end_date - start_date).days + 1
        return start_date.isoformat(), end_date.isoformat(), duration

    match = SINGLE_DATE_RE.search(raw_value)
    if not match:
        return "", "", None

    year, month, day = match.groups()
    only_date = date(int(year), int(month), int(day)).isoformat()
    return only_date, only_date, 1


def extract_checkbox_value(runs: list[dict[str, str]], label: str) -> str:
    for index, item in enumerate(runs):
        if item["type"] != "text" or label not in item["value"]:
            continue

        symbol_values: list[str] = []
        for candidate in runs[index + 1 :]:
            if candidate["type"] == "text" and any(
                marker in candidate["value"]
                for marker in (
                    "一、",
                    "二、",
                    "三、",
                    "会投工程部意见：",
                    "分公司工程部意见：",
                )
            ):
                break
            if candidate["type"] == "sym":
                symbol_values.append(candidate["value"])
                if len(symbol_values) == 2:
                    break

        if len(symbol_values) < 2:
            return ""
        if symbol_values[0] == CHECKED_SYMBOL:
            return "是"
        if symbol_values[1] == CHECKED_SYMBOL:
            return "否"
        return ""
    return ""


def next_symbol_value(runs: list[dict[str, str]], start_index: int) -> str:
    for candidate in runs[start_index + 1 :]:
        if candidate["type"] == "sym":
            return candidate["value"]
        if candidate["type"] == "text" and any(
            option in candidate["value"] for option in WORK_TYPES
        ):
            return ""
    return ""


def extract_dangerous_work(runs: list[dict[str, str]]) -> str:
    selected: list[str] = []
    cursor = 0

    for option in WORK_TYPES:
        while cursor < len(runs):
            item = runs[cursor]
            cursor += 1
            if item["type"] != "text" or option not in item["value"]:
                continue
            if next_symbol_value(runs, cursor - 1) == CHECKED_SYMBOL:
                selected.append(option)
            break

    return "、".join(selected)


def detect_attachment_dir(docx_path: Path) -> Path | None:
    candidates = [
        docx_path.with_name(f"{docx_path.stem}_附件"),
        docx_path.with_name(f"{docx_path.stem}附件"),
        docx_path.with_name(docx_path.stem),
    ]
    for candidate in candidates:
        if candidate.is_dir():
            return candidate
    return None


def collect_jpg_files(directory: Path | None) -> tuple[Path, ...]:
    if not directory or not directory.is_dir():
        return ()
    return tuple(
        sorted(
            (
                path
                for path in directory.iterdir()
                if path.is_file() and path.suffix.lower() in JPG_SUFFIXES
            ),
            key=lambda path: path.name.lower(),
        )
    )


def collect_pdf_files(
    input_dir: Path, skipped_pdf_names: set[str] | None = None
) -> list[Path]:
    skipped_pdf_names = skipped_pdf_names or set()
    return sorted(
        (
            path
            for path in input_dir.rglob("*")
            if path.is_file() and path.suffix.lower() == PDF_SUFFIX
            and path.name not in skipped_pdf_names
        ),
        key=lambda path: str(path).lower(),
    )


def build_pdf_path_index(
    input_dir: Path, skipped_pdf_names: set[str] | None = None
) -> dict[str, Path]:
    return {
        pdf_path.name: pdf_path
        for pdf_path in collect_pdf_files(input_dir, skipped_pdf_names)
    }


def read_pdf_plain_text(pdf_path: Path) -> str:
    document = fitz.open(pdf_path)
    page_texts: list[str] = []
    try:
        for page in document:
            page_texts.append(str(page.get_text("text") or ""))
    finally:
        document.close()
    return "\n".join(page_texts)


def read_pdf_ai_ocr_text(pdf_path: Path, client: LlamaCppClient) -> str:
    document = fitz.open(pdf_path)
    page_texts: list[str] = []
    try:
        page_count = min(len(document), OCR_PAGE_LIMIT)
        for page_index in range(page_count):
            page = document[page_index]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            page_texts.append(
                client.extract_image_bytes_text(
                    pixmap.tobytes("png"),
                    "image/png",
                    PDF_PAGE_TEXT_PROMPT,
                    max_tokens=768,
                )
            )
    finally:
        document.close()
    return "\n".join(page_texts)


def pdf_content_fingerprint(pdf_path: Path) -> str:
    digest = hashlib.sha256()
    with pdf_path.open("rb") as pdf_file:
        while chunk := pdf_file.read(PDF_FINGERPRINT_CHUNK_SIZE):
            digest.update(chunk)
    return digest.hexdigest()


def load_pdf_text_cache(cache_path: Path) -> dict[str, dict[str, Any]]:
    if not cache_path.is_file():
        return {}

    try:
        payload = json.loads(cache_path.read_text(encoding="utf-8"))
        if not isinstance(payload, dict):
            raise ValueError("缓存根节点不是对象")
        if payload.get("version") != PDF_TEXT_CACHE_VERSION:
            LOGGER.info(
                "PDF 文本缓存版本为 %s，当前版本为 %s，忽略旧缓存",
                payload.get("version") or "未标记",
                PDF_TEXT_CACHE_VERSION,
            )
            return {}
        raw_entries = payload.get("entries")
        if not isinstance(raw_entries, dict):
            raise ValueError("entries 不是对象")
        entries = {
            fingerprint: entry
            for fingerprint, entry in raw_entries.items()
            if isinstance(fingerprint, str)
            and isinstance(entry, dict)
            and isinstance(entry.get("text"), str)
        }
        LOGGER.info("已读取 %s 条 PDF 文本特征缓存: %s", len(entries), cache_path)
        return entries
    except (OSError, ValueError, TypeError, json.JSONDecodeError) as exc:
        LOGGER.warning("PDF 文本缓存读取失败，将重新识别: %s (%s)", cache_path, exc)
        return {}


def save_pdf_text_cache(
    cache_path: Path, entries: dict[str, dict[str, Any]]
) -> None:
    payload = {
        "version": PDF_TEXT_CACHE_VERSION,
        "fingerprint": "sha256",
        "entries": entries,
    }
    temporary_path = cache_path.with_name(f"{cache_path.name}.tmp")
    try:
        temporary_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        os.replace(temporary_path, cache_path)
    except OSError as exc:
        LOGGER.warning("PDF 文本缓存写入失败: %s (%s)", cache_path, exc)


def cache_pdf_text(
    cache_path: Path,
    entries: dict[str, dict[str, Any]],
    fingerprint: str,
    pdf_path: Path,
    text: str,
    method: str,
) -> None:
    entries[fingerprint] = {
        "text": text,
        "method": method,
        "file_name": pdf_path.name,
        "file_size": pdf_path.stat().st_size,
        "cached_at": datetime.now().isoformat(timespec="seconds"),
    }
    save_pdf_text_cache(cache_path, entries)


def get_cached_pdf_text(
    entries: dict[str, dict[str, Any]], fingerprint: str
) -> tuple[str | None, str]:
    entry = entries.get(fingerprint)
    if not entry or not isinstance(entry.get("text"), str):
        return None, ""
    return entry["text"], str(entry.get("method") or "unknown")


def read_pdf_text(pdf_path: Path, client: LlamaCppClient | None) -> str:
    plain_text = read_pdf_plain_text(pdf_path)
    if count_cjk_chars(plain_text) >= MIN_PLAIN_PDF_CJK_CHARS:
        return normalize_match_text(plain_text)
    if client is None:
        return normalize_match_text(plain_text)
    return normalize_match_text(read_pdf_ai_ocr_text(pdf_path, client))


def is_target_pdf_name(pdf_name: str) -> bool:
    return PDF_TARGET_NAME_RE.fullmatch(normalize_digits(pdf_name)) is not None


def build_pdf_target_path(pdf_path: Path, application_no: str) -> Path:
    return pdf_path.with_name(f"{PDF_TARGET_NAME_PREFIX}{application_no}.pdf")


def extract_pdf_rename_application_no(pdf_text: str) -> str:
    normalized = normalize_digits(pdf_text or "")
    compact_text = re.sub(r"\s+", "", normalized)
    if normalize_match_text("工程类-主体质保施工") not in normalize_match_text(normalized):
        return ""

    match = PDF_APPLICATION_NO_RE.search(compact_text)
    return match.group(1) if match else ""


def rename_pdf_by_application_no(pdf_path: Path, application_no: str) -> bool:
    target_path = build_pdf_target_path(pdf_path, application_no)
    if pdf_path.resolve() == target_path.resolve():
        LOGGER.info("PDF 已符合命名规则，跳过: %s", pdf_path.name)
        return False
    if target_path.exists():
        pdf_path.unlink()
        LOGGER.warning(
            "AI 识别到重名 PDF，保留已有文件并删除多余文件: 保留=%s，删除=%s，申请编号=%s",
            target_path.name,
            pdf_path.name,
            application_no,
        )
        return False

    pdf_path.rename(target_path)
    target_path.touch()
    LOGGER.info("PDF 重命名成功: %s -> %s", pdf_path.name, target_path.name)
    return True


def rename_subject_warranty_pdfs_by_local_ai(
    input_dir: Path,
    repo_root: Path,
    skipped_pdf_names: set[str] | None = None,
) -> int:
    skipped_pdf_names = skipped_pdf_names or set()
    pdf_files = [
        pdf_path
        for pdf_path in collect_pdf_files(input_dir, skipped_pdf_names)
        if not is_target_pdf_name(pdf_path.name)
    ]
    if not pdf_files:
        LOGGER.info("未发现需要本地 AI 识别重命名的 PDF")
        return 0

    text_cache_path = input_dir / PDF_TEXT_CACHE_NAME
    text_cache = load_pdf_text_cache(text_cache_path)
    config = LlamaCppConfig.from_repo(repo_root)
    client = LlamaCppClient(config, repo_root)
    client_ready = False
    renamed_count = 0
    try:
        LOGGER.info("发现 %s 个待识别 PDF，开始调用本地 AI", len(pdf_files))

        for index, pdf_path in enumerate(pdf_files, start=1):
            LOGGER.info("识别 PDF %s/%s: %s", index, len(pdf_files), pdf_path.name)
            try:
                fingerprint = pdf_content_fingerprint(pdf_path)
                pdf_text, cache_method = get_cached_pdf_text(
                    text_cache, fingerprint
                )
                if pdf_text is not None:
                    LOGGER.info(
                        "复用 PDF 文本特征缓存: %s (sha256=%s, 来源=%s)",
                        pdf_path.name,
                        fingerprint[:12],
                        cache_method,
                    )
                else:
                    if not client_ready:
                        if not config.mmproj_path and not client.is_server_available():
                            LOGGER.warning(
                                "未配置 LLAMACPP_MMPROJ_PATH，且本地 AI 服务未运行，"
                                "跳过剩余 PDF 识别重命名"
                            )
                            break
                        client.ensure_server()
                        client.assert_model_available()
                        client_ready = True
                    pdf_text = normalize_match_text(
                        read_pdf_ai_ocr_text(pdf_path, client)
                    )
                    cache_pdf_text(
                        text_cache_path,
                        text_cache,
                        fingerprint,
                        pdf_path,
                        pdf_text,
                        "ocr",
                    )
                application_no = extract_pdf_rename_application_no(pdf_text)
                if not application_no:
                    LOGGER.info("PDF 未匹配主体质保施工申请编号规则: %s", pdf_path.name)
                    continue
                if rename_pdf_by_application_no(pdf_path, application_no):
                    renamed_count += 1
            except Exception as exc:
                LOGGER.warning("PDF 识别重命名失败，已跳过: %s (%s)", pdf_path.name, exc)
    except Exception as exc:
        LOGGER.warning("本地 AI PDF 重命名不可用，已跳过剩余 PDF: %s", exc)
    finally:
        client.shutdown_server()
    return renamed_count


def build_pdf_text_index(
    input_dir: Path,
    excluded_pdf_names: set[str] | None = None,
    skipped_pdf_names: set[str] | None = None,
    repo_root: Path | None = None,
) -> dict[Path, str]:
    pdf_texts: dict[Path, str] = {}
    ai_client: LlamaCppClient | None = None
    excluded_pdf_names = excluded_pdf_names or set()
    skipped_pdf_names = skipped_pdf_names or set()
    repo_root = repo_root or Path(__file__).resolve().parent
    text_cache_path = input_dir / PDF_TEXT_CACHE_NAME
    text_cache = load_pdf_text_cache(text_cache_path)
    cache_hit_count = 0
    cache_write_count = 0
    pdf_files = [
        pdf_path
        for pdf_path in collect_pdf_files(input_dir, skipped_pdf_names)
        if pdf_path.name not in excluded_pdf_names
    ]
    if skipped_pdf_names:
        LOGGER.info("跳过 %s 个 env 配置排除的 PDF", len(skipped_pdf_names))
    if excluded_pdf_names:
        LOGGER.info("跳过 %s 个已有匹配结果的 PDF", len(excluded_pdf_names))
    LOGGER.info("发现 %s 个 PDF，开始建立匹配索引", len(pdf_files))
    try:
        for index, pdf_path in enumerate(pdf_files, start=1):
            LOGGER.info("读取 PDF %s/%s: %s", index, len(pdf_files), pdf_path.name)
            try:
                fingerprint = pdf_content_fingerprint(pdf_path)
                cached_text, cache_method = get_cached_pdf_text(
                    text_cache, fingerprint
                )
                if cached_text is not None:
                    pdf_texts[pdf_path] = cached_text
                    cache_hit_count += 1
                    LOGGER.info(
                        "复用 PDF 文本特征缓存: %s (sha256=%s, 来源=%s)",
                        pdf_path.name,
                        fingerprint[:12],
                        cache_method,
                    )
                    continue

                plain_text = read_pdf_plain_text(pdf_path)
                if count_cjk_chars(plain_text) >= MIN_PLAIN_PDF_CJK_CHARS:
                    pdf_texts[pdf_path] = normalize_match_text(plain_text)
                    cache_pdf_text(
                        text_cache_path,
                        text_cache,
                        fingerprint,
                        pdf_path,
                        pdf_texts[pdf_path],
                        "plain",
                    )
                    cache_write_count += 1
                    LOGGER.info("PDF 文本提取完成: %s", pdf_path.name)
                    continue

                if ai_client is None:
                    LOGGER.info("普通文本不可用，初始化本地 AI PDF OCR")
                    ai_client = LlamaCppClient(
                        LlamaCppConfig.from_repo(repo_root),
                        repo_root,
                    )
                    ai_client.ensure_server()
                    ai_client.assert_model_available()
                LOGGER.info("开始本地 AI OCR 识别 PDF: %s", pdf_path.name)
                pdf_texts[pdf_path] = normalize_match_text(
                    read_pdf_ai_ocr_text(pdf_path, ai_client)
                )
                cache_pdf_text(
                    text_cache_path,
                    text_cache,
                    fingerprint,
                    pdf_path,
                    pdf_texts[pdf_path],
                    "ocr",
                )
                cache_write_count += 1
                LOGGER.info("本地 AI OCR 识别完成: %s", pdf_path.name)
            except Exception as exc:
                pdf_texts[pdf_path] = ""
                LOGGER.warning("PDF 读取失败，已跳过: %s (%s)", pdf_path, exc)
    finally:
        if ai_client is not None:
            ai_client.shutdown_server()
    LOGGER.info(
        "PDF 文本特征缓存统计: 命中 %s，新增 %s",
        cache_hit_count,
        cache_write_count,
    )
    return pdf_texts


PDF_DATE_TOKEN_PATTERN = (
    r"(?<!\d)(\d{4})\s*(?:年|[./．/~\-—–－])\s*(\d{1,2})"
    r"\s*(?:月|[./．/~\-—–－])\s*(\d{1,2})\s*(?:日)?(?!\d)"
)
PDF_DATE_TOKEN_RE = re.compile(PDF_DATE_TOKEN_PATTERN)
PDF_COMPACT_DATE_RE = re.compile(r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)")
PDF_START_DATE_RE = re.compile(
    rf"施工开始(?:时间|日期)?\s*[:：]?\s*{PDF_DATE_TOKEN_PATTERN}"
)
PDF_END_DATE_RE = re.compile(
    rf"施工结束(?:时间|日期)?\s*[:：]?\s*{PDF_DATE_TOKEN_PATTERN}"
)


def date_match_groups_to_iso(groups: tuple[str, str, str]) -> str:
    try:
        return date(*(int(value) for value in groups)).isoformat()
    except ValueError:
        return ""


def normalize_date_for_pdf_match(value: str) -> str:
    normalized = normalize_digits(str(value or ""))
    match = PDF_DATE_TOKEN_RE.search(normalized) or PDF_COMPACT_DATE_RE.search(
        normalized
    )
    return date_match_groups_to_iso(match.groups()) if match else ""


def extract_pdf_dates(pdf_text: str) -> set[str]:
    normalized = normalize_digits(pdf_text or "")
    dates = {
        date_match_groups_to_iso(match.groups())
        for match in PDF_DATE_TOKEN_RE.finditer(normalized)
    }
    dates.update(
        date_match_groups_to_iso(match.groups())
        for match in PDF_COMPACT_DATE_RE.finditer(normalized)
    )
    dates.discard("")
    return dates


def pdf_construction_dates_match(
    pdf_text: str, construction_start_date: str, construction_end_date: str
) -> bool:
    expected_start = normalize_date_for_pdf_match(construction_start_date)
    expected_end = normalize_date_for_pdf_match(construction_end_date)
    if not expected_start or not expected_end:
        return False

    normalized_pdf_text = normalize_digits(pdf_text or "")
    start_match = PDF_START_DATE_RE.search(normalized_pdf_text)
    end_match = PDF_END_DATE_RE.search(normalized_pdf_text)
    if start_match and end_match:
        actual_start = date_match_groups_to_iso(start_match.groups())
        actual_end = date_match_groups_to_iso(end_match.groups())
        return actual_start == expected_start and actual_end == expected_end

    # OCR 偶尔会漏掉字段标签；此时仍要求 PDF 中同时出现申请单的起止日期。
    pdf_dates = extract_pdf_dates(normalized_pdf_text)
    return expected_start in pdf_dates and expected_end in pdf_dates


def find_matching_pdf_paths(
    construction_area: str,
    work_content: str,
    pdf_texts: dict[Path, str],
    construction_start_date: str = "",
    construction_end_date: str = "",
) -> list[Path]:
    area_key = normalize_match_text(construction_area)
    content_key = normalize_match_text(work_content)
    if not area_key or not content_key:
        return []

    content_keys = [content_key]
    if content_key.startswith(area_key):
        short_content_key = content_key[len(area_key) :]
        if short_content_key:
            content_keys.append(short_content_key)

    matched_paths = [
        pdf_path
        for pdf_path, pdf_text in pdf_texts.items()
        if area_key in pdf_text and any(key in pdf_text for key in content_keys)
        and pdf_construction_dates_match(
            pdf_text, construction_start_date, construction_end_date
        )
    ]
    return matched_paths


def find_matching_pdf_names(
    construction_area: str,
    work_content: str,
    pdf_texts: dict[Path, str],
    construction_start_date: str = "",
    construction_end_date: str = "",
) -> str:
    matched_names = [
        pdf_path.name
        for pdf_path in find_matching_pdf_paths(
            construction_area,
            work_content,
            pdf_texts,
            construction_start_date,
            construction_end_date,
        )
    ]
    return "；".join(matched_names)


def detect_application_image(
    current_docx_path: Path,
    original_docx_path: Path,
    attachment_dir: Path | None,
) -> Path | None:
    candidates: list[Path] = []
    for docx_path in (current_docx_path, original_docx_path):
        for suffix in JPG_SUFFIXES:
            candidates.append(docx_path.with_suffix(suffix))

    for candidate in candidates:
        if candidate.is_file():
            return candidate

    jpg_files = collect_jpg_files(attachment_dir)
    return jpg_files[0] if jpg_files else None


def build_target_name(start_date: str, work_content: str) -> str:
    if not start_date or not work_content:
        raise ValueError("缺少施工开始时间或施工内容，无法生成新文件名")
    return f"{start_date}_{sanitize_filename_part(work_content)}_质保作业申请单.docx"


def parse_document(docx_path: Path) -> dict[str, Any]:
    form_text = normalize_whitespace(read_first_form_text(docx_path))
    runs = read_first_form_runs(docx_path)

    date_range_raw = capture_between(form_text, "施工日期：", "施工内容：")
    start_date, end_date, duration_days = parse_date_range(date_range_raw)
    constructor_section = capture_between(
        form_text, "施工内容：", "一、影响、改动消防设备设施"
    )
    phone_match = re.search(r"联系电话：\s*(\d+)", constructor_section)

    return {
        "项目名称": capture_between(form_text, "项目名称：", "质保单位（盖章）："),
        "质保单位": capture_between(form_text, "质保单位（盖章）：", "分包单位："),
        "分包单位": capture_between(form_text, "分包单位：", "质保负责人："),
        "质保负责人": capture_between(form_text, "质保负责人：", "联系电话："),
        "质保负责人联系电话": capture_between(form_text, "联系电话：", "施工区域："),
        "施工区域": capture_between(form_text, "施工区域：", "施工日期："),
        "施工开始时间": start_date,
        "施工结束时间": end_date,
        "时长天": duration_days,
        "施工内容": capture_between(form_text, "施工内容：", "施工负责人："),
        "施工负责人": capture_between(form_text, "施工负责人：", "联系电话："),
        "施工负责人联系电话": phone_match.group(1) if phone_match else "",
        "影响改动消防设备设施": extract_checkbox_value(
            runs, "一、影响、改动消防设备设施"
        ),
        "影响堵塞应急疏散通道": extract_checkbox_value(
            runs, "二、影响、堵塞应急疏散通道"
        ),
        "危险作业": extract_dangerous_work(runs),
    }


def export_excel(records: list[Record], output_path: Path) -> Path:
    workbook = Workbook()
    summary_sheet = workbook.active
    if summary_sheet is None:
        raise RuntimeError("Excel 工作簿未创建默认工作表")
    summary_sheet.title = "汇总"

    headers = [
        "项目名称",
        "质保单位",
        "分包单位",
        "质保负责人",
        "质保负责人联系电话",
        "施工区域",
        "施工开始时间",
        "施工结束时间",
        "时长（天）",
        "施工内容",
        "施工负责人",
        "施工负责人联系电话",
        "影响、改动消防设备设施",
        "影响、堵塞应急疏散通道",
        "危险作业",
        "申请单文件",
        "申请单附图",
        "附件目录",
        "匹配PDF文件名",
        "原文件名",
        "新文件名",
        "文件路径",
        "处理状态",
    ]
    summary_sheet.append(headers)

    for record in records:
        summary_sheet.append(
            [
                record.项目名称,
                record.质保单位,
                record.分包单位,
                record.质保负责人,
                record.质保负责人联系电话,
                record.施工区域,
                record.施工开始时间,
                record.施工结束时间,
                record.时长天,
                record.施工内容,
                record.施工负责人,
                record.施工负责人联系电话,
                record.影响改动消防设备设施,
                record.影响堵塞应急疏散通道,
                record.危险作业,
                "打开文件" if record.申请单文件链接 else "",
                Path(record.申请单附图链接).name if record.申请单附图链接 else "",
                "打开目录" if record.附件目录链接 else "",
                record.匹配PDF文件名,
                record.原文件名,
                record.新文件名,
                record.文件路径,
                record.处理状态,
            ]
        )

    apply_summary_styles(summary_sheet)
    fill_summary_hyperlinks(summary_sheet, records)

    note_sheet = workbook.create_sheet("说明")
    build_note_sheet(note_sheet, records, output_path.name)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        workbook.save(output_path)
        return output_path
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        fallback_path = output_path.with_name(
            f"{output_path.stem}_{timestamp}{output_path.suffix}"
        )
        workbook.save(fallback_path)
        return fallback_path


def apply_summary_styles(summary_sheet: Any) -> None:
    thin_black = Side(style="thin", color="000000")
    border = Border(
        top=thin_black,
        bottom=thin_black,
        left=thin_black,
        right=thin_black,
    )
    header_fill = PatternFill(fill_type="solid", fgColor="5B6F84")
    band_fill = PatternFill(fill_type="solid", fgColor="D9E5F3")
    plain_fill = PatternFill(fill_type=None)
    header_font = Font(bold=True, color="FFFFFF")
    hyperlink_font = Font(color="0563C1", underline="single")
    centered = Alignment(
        horizontal="center",
        vertical="center",
        wrap_text=True,
    )
    wrapped = Alignment(vertical="center", wrap_text=True)

    base_column_widths = {
        1: 16,
        2: 16,
        3: 16,
        4: 12,
        5: 16,
        6: 20,
        7: 12,
        8: 12,
        9: 10,
        10: 24,
        11: 12,
        12: 16,
        13: 16,
        14: 18,
        15: 24,
        16: 12,
        17: 20,
    }
    header_widths = {
        "附件目录": 12,
        "匹配PDF文件名": 32,
        "原文件名": 28,
        "新文件名": 32,
        "文件路径": 40,
        "处理状态": 12,
    }
    for col_idx, width in base_column_widths.items():
        summary_sheet.column_dimensions[get_column_letter(col_idx)].width = width
    for col_idx in range(17, summary_sheet.max_column + 1):
        header = summary_sheet.cell(row=1, column=col_idx).value
        width = base_column_widths.get(col_idx) or header_widths.get(header, 12)
        summary_sheet.column_dimensions[get_column_letter(col_idx)].width = width

    summary_sheet.freeze_panes = "A2"
    summary_sheet.auto_filter.ref = (
        f"A1:{get_column_letter(summary_sheet.max_column)}{summary_sheet.max_row}"
    )
    summary_sheet.sheet_view.showGridLines = False
    summary_sheet.row_dimensions[1].height = 24

    for cell in summary_sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border
        cell.alignment = centered

    for row_idx, row in enumerate(
        summary_sheet.iter_rows(
            min_row=2,
            max_row=summary_sheet.max_row,
            min_col=1,
            max_col=summary_sheet.max_column,
        ),
        start=2,
    ):
        if not any(cell.value not in (None, "") for cell in row):
            continue
        fill = band_fill if row_idx % 2 == 0 else plain_fill
        for cell in row:
            cell.alignment = wrapped
            cell.border = border
            cell.fill = fill
        for cell in row:
            header = summary_sheet.cell(row=1, column=cell.column).value
            if header in {
                "申请单文件",
                "申请单附图",
                "附件目录",
                "匹配PDF文件名",
            }:
                cell.font = hyperlink_font

    for row_idx in range(2, summary_sheet.max_row + 1):
        summary_sheet[f"G{row_idx}"].number_format = "yyyy-mm-dd"
        summary_sheet[f"H{row_idx}"].number_format = "yyyy-mm-dd"
        summary_sheet[f"I{row_idx}"].number_format = "0"
        summary_sheet.row_dimensions[row_idx].height = 42


def fill_summary_hyperlinks(summary_sheet: Any, records: list[Record]) -> None:
    header_columns = {
        summary_sheet.cell(row=1, column=col_idx).value: col_idx
        for col_idx in range(1, summary_sheet.max_column + 1)
    }
    docx_col = header_columns.get("申请单文件")
    image_col = header_columns.get("申请单附图")
    attachment_col = header_columns.get("附件目录")
    pdf_col = header_columns.get("匹配PDF文件名")
    for row_idx, record in enumerate(records, start=2):
        if record.申请单文件链接 and docx_col:
            summary_sheet.cell(row=row_idx, column=docx_col).hyperlink = (
                record.申请单文件链接
            )
        if record.申请单附图链接 and image_col:
            summary_sheet.cell(row=row_idx, column=image_col).hyperlink = (
                record.申请单附图链接
            )
        if record.附件目录链接 and attachment_col:
            summary_sheet.cell(row=row_idx, column=attachment_col).hyperlink = (
                record.附件目录链接
            )
        if record.匹配PDF文件链接 and pdf_col:
            summary_sheet.cell(row=row_idx, column=pdf_col).hyperlink = (
                record.匹配PDF文件链接
            )


def build_note_sheet(
    note_sheet: Any, records: list[Record], workbook_name: str
) -> None:
    rows = [
        ["项目", "说明"],
        ["生成时间", date.today().isoformat()],
        ["文档总数", len(records)],
        ["重命名数量", sum(1 for record in records if record.处理状态 == "已重命名")],
        ["汇总文件", workbook_name],
        ["PDF匹配规则版本", PDF_MATCH_CACHE_VERSION],
        [
            "说明",
            "申请单文件、申请单附图和附件目录列可直接点击打开本地文件或目录；匹配PDF文件名列按施工区域、施工内容、施工开始时间和施工结束时间匹配。",
        ],
    ]
    for row in rows:
        note_sheet.append(row)

    thin_black = Side(style="thin", color="000000")
    border = Border(
        top=thin_black, bottom=thin_black, left=thin_black, right=thin_black
    )
    header_fill = PatternFill(fill_type="solid", fgColor="5B6F84")
    wrapped = Alignment(vertical="center", wrap_text=True)

    note_sheet.column_dimensions["A"].width = 18
    note_sheet.column_dimensions["B"].width = 40

    for cell in note_sheet[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF")
        cell.border = border
        cell.alignment = wrapped

    for row in note_sheet.iter_rows(
        min_row=2, max_row=note_sheet.max_row, min_col=1, max_col=2
    ):
        for cell in row:
            cell.border = border
            cell.alignment = wrapped


def collect_docx_files(input_dir: Path) -> list[Path]:
    return sorted(path for path in input_dir.glob("*.docx") if path.is_file())


def collect_ai_candidate_images(input_dir: Path) -> list[Path]:
    return sorted(
        (
            path
            for path in input_dir.glob("*")
            if path.is_file()
            and path.suffix.lower() in AI_IMAGE_SUFFIXES
            and not DATE_PREFIX_RE.match(path.name)
        ),
        key=lambda path: path.name.lower(),
    )


def word_content_from_name(word_name: str) -> str:
    match = WORD_CONTENT_NAME_RE.fullmatch(word_name)
    return match.group(2) if match else ""


def word_date_from_name(word_name: str) -> date | None:
    match = WORD_CONTENT_NAME_RE.fullmatch(word_name)
    if not match:
        return None
    try:
        return date.fromisoformat(match.group(1))
    except ValueError:
        return None


def build_recent_word_image_match_index(
    records: list[Record],
    run_date: date | None = None,
    lookback_days: int = 2,
    lookahead_days: int = 14,
) -> list[tuple[Record, str, str]]:
    run_date = run_date or date.today()
    earliest_date = run_date - timedelta(days=lookback_days)
    latest_date = run_date + timedelta(days=lookahead_days)
    candidates: list[tuple[Record, str, str]] = []
    for record in records:
        word_date = word_date_from_name(record.新文件名)
        if word_date is None or word_date < earliest_date or word_date > latest_date:
            continue
        content = word_content_from_name(record.新文件名)
        content_key = normalize_match_text(content)
        if content_key:
            candidates.append((record, content, content_key))
    return candidates


def find_best_word_image_match(
    image_text: str, candidates: list[tuple[Record, str, str]]
) -> tuple[Record, str] | None:
    text_key = normalize_match_text(image_text)
    if not text_key:
        return None
    matches = [
        (record, content)
        for record, content, content_key in candidates
        if content_key and content_key in text_key
    ]
    if not matches:
        return None
    return max(matches, key=lambda item: len(normalize_match_text(item[1])))


def is_worker_list_image_text(image_text: str) -> bool:
    text_key = normalize_match_text(image_text)
    return all(field in text_key for field in ("姓名", "电话", "性别"))


def build_worker_list_image_target(image_path: Path) -> Path:
    if image_path.stem.endswith(WORKER_LIST_IMAGE_SUFFIX):
        return image_path
    return image_path.with_name(
        f"{image_path.stem}{WORKER_LIST_IMAGE_SUFFIX}{image_path.suffix}"
    )


def rename_image_file(image_path: Path, target_path: Path, log_label: str) -> bool:
    if target_path == image_path:
        LOGGER.info("图片已符合%s命名规则: %s", log_label, image_path.name)
        return False
    if target_path.exists() and target_path.resolve() != image_path.resolve():
        LOGGER.warning(
            "图片目标文件已存在，跳过重命名: %s -> %s",
            image_path.name,
            target_path.name,
        )
        return False

    image_path.rename(target_path)
    target_path.touch()
    return True


def rename_matched_images_by_local_ai(
    input_dir: Path, records: list[Record], repo_root: Path
) -> int:
    images = collect_ai_candidate_images(input_dir)
    candidates = build_recent_word_image_match_index(records)
    if not images:
        LOGGER.info("未发现需要本地 AI 识别重命名的图片")
        return 0
    if candidates:
        LOGGER.info("图片匹配 Word 候选限定为运行日前 2 天至后 14 天内: %s 个", len(candidates))
    else:
        LOGGER.info("未发现运行日前 2 天至后 14 天内可用于图片匹配的 Word 文件名")

    config = LlamaCppConfig.from_repo(repo_root)
    client = LlamaCppClient(config, repo_root)
    renamed_count = 0
    try:
        if not config.mmproj_path and not client.is_server_available():
            LOGGER.warning(
                "未配置 LLAMACPP_MMPROJ_PATH，且本地 AI 服务未运行，跳过图片识别重命名"
            )
            return 0
        LOGGER.info("发现 %s 个待识别图片，开始调用本地 AI", len(images))
        client.ensure_server()
        client.assert_model_available()

        for index, image_path in enumerate(images, start=1):
            LOGGER.info("识别图片 %s/%s: %s", index, len(images), image_path.name)
            try:
                image_text = client.extract_image_text(image_path)
            except Exception as exc:
                LOGGER.warning("图片识别失败，已跳过: %s (%s)", image_path.name, exc)
                continue

            is_worker_list = is_worker_list_image_text(image_text)
            match = find_best_word_image_match(image_text, candidates)
            if not match:
                if is_worker_list:
                    target_path = build_worker_list_image_target(image_path)
                    if rename_image_file(image_path, target_path, "工人名单"):
                        LOGGER.info(
                            "图片识别为工人名单: %s -> %s",
                            image_path.name,
                            target_path.name,
                        )
                        renamed_count += 1
                    continue
                LOGGER.info("图片未匹配到 Word 文件名内容: %s", image_path.name)
                continue

            record, content = match
            target_stem = Path(record.新文件名).stem
            target_suffix = image_path.suffix.lower()
            if is_worker_list:
                target_stem = f"{target_stem}{WORKER_LIST_IMAGE_SUFFIX}"
                target_suffix = image_path.suffix
            target_path = input_dir / f"{target_stem}{target_suffix}"
            if not rename_image_file(image_path, target_path, "Word 匹配图片"):
                continue

            record.申请单附图链接 = str(target_path)
            LOGGER.info("图片匹配成功: %s -> %s (%s)", image_path.name, target_path.name, content)
            renamed_count += 1
    except Exception as exc:
        LOGGER.warning("本地 AI 图片重命名不可用，已跳过剩余图片: %s", exc)
    finally:
        client.shutdown_server()
    return renamed_count


def pdf_match_cache_key(
    construction_area: str,
    work_content: str,
    construction_start_date: str = "",
    construction_end_date: str = "",
) -> str:
    return (
        f"{normalize_match_text(construction_start_date)}|"
        f"{normalize_match_text(construction_end_date)}|"
        f"{normalize_match_text(construction_area)}|"
        f"{normalize_match_text(work_content)}"
    )


def normalize_excel_cache_value(value: Any) -> str:
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    return normalize_whitespace(str(value or ""))


def collect_cached_pdf_names(pdf_match_cache: dict[str, str]) -> set[str]:
    cached_pdf_names: set[str] = set()
    for pdf_names in pdf_match_cache.values():
        cached_pdf_names.update(split_pdf_names(pdf_names))
    return cached_pdf_names


def load_existing_pdf_match_cache(excel_path: Path) -> dict[str, str]:
    if not excel_path.is_file():
        LOGGER.info("未找到已有汇总表，PDF 匹配缓存为空: %s", excel_path)
        return {}

    LOGGER.info("读取已有 PDF 匹配缓存: %s", excel_path)
    workbook = load_workbook(excel_path, read_only=False, data_only=True)
    try:
        if "说明" not in workbook.sheetnames:
            LOGGER.info("已有汇总表没有 PDF 匹配规则版本，忽略旧缓存")
            return {}
        note_sheet = workbook["说明"]
        cache_version = next(
            (
                normalize_whitespace(str(row[1].value or ""))
                for row in note_sheet.iter_rows(min_col=1, max_col=2)
                if row[0].value == "PDF匹配规则版本"
            ),
            "",
        )
        if cache_version != PDF_MATCH_CACHE_VERSION:
            LOGGER.info(
                "已有汇总表 PDF 匹配规则版本为 %s，当前版本为 %s，忽略旧缓存",
                cache_version or "未标记",
                PDF_MATCH_CACHE_VERSION,
            )
            return {}

        sheet = workbook["汇总"] if "汇总" in workbook.sheetnames else workbook.active
        if sheet is None:
            LOGGER.info("已有汇总表没有可读取的工作表")
            return {}
        headers = {
            sheet.cell(row=1, column=col_idx).value: col_idx
            for col_idx in range(1, sheet.max_column + 1)
        }
        area_col = headers.get("施工区域")
        start_date_col = headers.get("施工开始时间")
        end_date_col = headers.get("施工结束时间")
        content_col = headers.get("施工内容")
        pdf_col = headers.get("匹配PDF文件名")
        if (
            not area_col
            or not start_date_col
            or not end_date_col
            or not content_col
            or not pdf_col
        ):
            LOGGER.info("已有汇总表缺少匹配缓存所需列")
            return {}

        cache: dict[str, str] = {}
        for row_idx in range(2, sheet.max_row + 1):
            pdf_names = normalize_whitespace(
                str(sheet.cell(row=row_idx, column=pdf_col).value or "")
            )
            if not pdf_names:
                continue
            cache[
                pdf_match_cache_key(
                    normalize_excel_cache_value(
                        sheet.cell(row=row_idx, column=area_col).value
                    ),
                    normalize_excel_cache_value(
                        sheet.cell(row=row_idx, column=content_col).value
                    ),
                    normalize_excel_cache_value(
                        sheet.cell(row=row_idx, column=start_date_col).value
                    ),
                    normalize_excel_cache_value(
                        sheet.cell(row=row_idx, column=end_date_col).value
                    ),
                )
            ] = pdf_names
        LOGGER.info("已读取 %s 条 PDF 匹配缓存", len(cache))
        return cache
    finally:
        workbook.close()


def resolve_pdf_link(pdf_names: str, pdf_path_index: dict[str, Path]) -> str:
    for pdf_name in split_pdf_names(pdf_names):
        pdf_path = pdf_path_index.get(pdf_name)
        if pdf_path:
            return str(pdf_path)
    return ""


def load_manual_pdf_matches(repo_root: Path) -> list[dict[str, str]]:
    env_values = load_env_file(repo_root / MANUAL_MATCH_ENV_NAME)
    raw_matches = env_values.get("MANUAL_PDF_MATCHES", "")
    if not raw_matches:
        return []

    parsed = yaml.safe_load(raw_matches)
    if not isinstance(parsed, list):
        raise ValueError(f"{MANUAL_MATCH_ENV_NAME} 中 MANUAL_PDF_MATCHES 必须是列表")

    matches: list[dict[str, str]] = []
    for item in parsed:
        if not isinstance(item, dict):
            raise ValueError(f"{MANUAL_MATCH_ENV_NAME} 中 MANUAL_PDF_MATCHES 条目必须是对象")
        application_image = normalize_whitespace(str(item.get("application_image") or ""))
        pdf = normalize_whitespace(str(item.get("pdf") or ""))
        if not application_image or not pdf:
            raise ValueError(
                f"{MANUAL_MATCH_ENV_NAME} 中每个特例必须包含 application_image 和 pdf"
            )
        matches.append({"application_image": application_image, "pdf": pdf})
    return matches


def build_file_name_index(input_dir: Path) -> dict[str, Path]:
    file_index: dict[str, Path] = {}
    for path in input_dir.rglob("*"):
        if path.is_file() and path.name not in file_index:
            file_index[path.name] = path
    return file_index


def resolve_manual_file_path(
    file_name: str,
    file_index: dict[str, Path],
    suffix_fallbacks: tuple[str, ...] = (),
) -> Path | None:
    matched_path = file_index.get(file_name)
    if matched_path:
        return matched_path

    source_path = Path(file_name)
    for suffix in suffix_fallbacks:
        fallback_name = f"{source_path.stem}{suffix}"
        matched_path = file_index.get(fallback_name)
        if matched_path:
            LOGGER.info("手工匹配文件名修正: %s -> %s", file_name, fallback_name)
            return matched_path
    return None


def normalize_manual_match_key(value: str) -> str:
    stem = Path(value).stem
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", normalize_digits(stem)).lower()


def find_manual_match_record(
    records: list[Record],
    application_image_name: str,
    application_image_path: Path | None,
) -> Record | None:
    image_names = [application_image_name]
    if application_image_path:
        image_names.append(application_image_path.name)

    image_keys = {normalize_manual_match_key(name) for name in image_names}
    for record in records:
        if record.申请单附图链接 and Path(record.申请单附图链接).name in image_names:
            return record

    for record in records:
        record_keys = {
            normalize_manual_match_key(record.新文件名),
            normalize_manual_match_key(record.原文件名),
        }
        if image_keys & record_keys:
            return record
    return None


def apply_manual_pdf_matches(
    records: list[Record],
    input_dir: Path,
    repo_root: Path,
) -> int:
    manual_matches = load_manual_pdf_matches(repo_root)
    if not manual_matches:
        LOGGER.info("未配置手工 PDF 匹配特例")
        return 0

    file_index = build_file_name_index(input_dir)
    applied_count = 0
    for item in manual_matches:
        image_name = item["application_image"]
        pdf_name = item["pdf"]
        image_path = resolve_manual_file_path(
            image_name, file_index, (".jpg", ".jpeg", ".png")
        )
        pdf_path = resolve_manual_file_path(pdf_name, file_index)
        if not pdf_path:
            LOGGER.warning("手工匹配 PDF 文件不存在，已跳过: %s", pdf_name)
            continue

        record = find_manual_match_record(records, image_name, image_path)
        if not record:
            LOGGER.warning("手工匹配未找到对应申请单记录，已跳过: %s", image_name)
            continue

        if image_path:
            record.申请单附图链接 = str(image_path)
        record.匹配PDF文件名 = pdf_path.name
        record.匹配PDF文件链接 = str(pdf_path)
        applied_count += 1
        LOGGER.info(
            "已应用手工 PDF 匹配: %s -> %s",
            image_path.name if image_path else image_name,
            pdf_path.name,
        )

    return applied_count


def process_documents(
    input_dir: Path,
    existing_excel_path: Path,
    skipped_pdf_names: set[str] | None = None,
    repo_root: Path | None = None,
) -> list[Record]:
    records: list[Record] = []
    skipped_pdf_names = skipped_pdf_names or set()
    repo_root = repo_root or Path(__file__).resolve().parent
    docx_files = collect_docx_files(input_dir)
    pdf_match_cache = load_existing_pdf_match_cache(existing_excel_path)
    pdf_path_index = build_pdf_path_index(input_dir, skipped_pdf_names)
    pdf_match_cache = normalize_pdf_match_cache_names(pdf_match_cache, pdf_path_index)
    cached_pdf_names = collect_cached_pdf_names(pdf_match_cache)
    pdf_texts: dict[Path, str] | None = None
    if skipped_pdf_names:
        LOGGER.info("env 配置排除 PDF: %s", "；".join(sorted(skipped_pdf_names)))
    if cached_pdf_names:
        LOGGER.info(
            "已有 Excel 匹配结果中的 PDF 将跳过重新识别: %s",
            "；".join(sorted(cached_pdf_names)),
        )

    LOGGER.info("发现 %s 个 Word 申请单，开始处理", len(docx_files))

    for doc_index, docx_path in enumerate(docx_files, start=1):
        LOGGER.info("处理申请单 %s/%s: %s", doc_index, len(docx_files), docx_path.name)
        original_name = docx_path.name
        parsed = parse_document(docx_path)
        target_name = build_target_name(parsed["施工开始时间"], parsed["施工内容"])

        current_path = docx_path
        status = "已符合命名规则"
        if not TARGET_NAME_RE.fullmatch(docx_path.name):
            target_path = docx_path.with_name(target_name)
            if target_path.exists() and target_path.resolve() != docx_path.resolve():
                raise FileExistsError(f"目标文件已存在，无法重命名: {target_path}")
            if target_path.name != docx_path.name:
                docx_path.rename(target_path)
                current_path = target_path
                current_path.touch()
                status = "已重命名"

        attachment_dir = detect_attachment_dir(current_path) or detect_attachment_dir(
            docx_path
        )
        application_image = detect_application_image(
            current_path, docx_path, attachment_dir
        )
        cache_key = pdf_match_cache_key(
            parsed["施工区域"],
            parsed["施工内容"],
            parsed["施工开始时间"],
            parsed["施工结束时间"],
        )
        matched_pdf_names = pdf_match_cache.get(cache_key, "")
        if matched_pdf_names:
            matched_pdf_link = resolve_pdf_link(matched_pdf_names, pdf_path_index)
            LOGGER.info("使用已有 PDF 匹配结果: %s -> %s", current_path.name, matched_pdf_names)
        else:
            if pdf_texts is None:
                pdf_texts = build_pdf_text_index(
                    input_dir,
                    cached_pdf_names,
                    skipped_pdf_names,
                    repo_root,
                )
            matched_pdf_paths = find_matching_pdf_paths(
                parsed["施工区域"],
                parsed["施工内容"],
                pdf_texts,
                parsed["施工开始时间"],
                parsed["施工结束时间"],
            )
            matched_pdf_names = PDF_MATCH_SEPARATOR.join(
                pdf_path.name for pdf_path in matched_pdf_paths
            )
            matched_pdf_link = str(matched_pdf_paths[0]) if matched_pdf_paths else ""
            if matched_pdf_names:
                LOGGER.info("匹配到 PDF: %s -> %s", current_path.name, matched_pdf_names)
            else:
                LOGGER.info("未匹配到 PDF: %s", current_path.name)
        records.append(
            Record(
                项目名称=parsed["项目名称"],
                质保单位=parsed["质保单位"],
                分包单位=parsed["分包单位"],
                质保负责人=parsed["质保负责人"],
                质保负责人联系电话=parsed["质保负责人联系电话"],
                施工区域=parsed["施工区域"],
                施工开始时间=parsed["施工开始时间"],
                施工结束时间=parsed["施工结束时间"],
                时长天=parsed["时长天"],
                施工内容=parsed["施工内容"],
                施工负责人=parsed["施工负责人"],
                施工负责人联系电话=parsed["施工负责人联系电话"],
                影响改动消防设备设施=parsed["影响改动消防设备设施"],
                影响堵塞应急疏散通道=parsed["影响堵塞应急疏散通道"],
                危险作业=parsed["危险作业"],
                原文件名=original_name,
                新文件名=current_path.name,
                文件路径=str(current_path),
                申请单文件链接=str(current_path),
                申请单附图链接=str(application_image) if application_image else "",
                附件目录=str(attachment_dir) if attachment_dir else "",
                附件目录链接=str(attachment_dir) if attachment_dir else "",
                匹配PDF文件名=matched_pdf_names,
                匹配PDF文件链接=matched_pdf_link,
                处理状态=status,
            )
        )

    return records


def resolve_input_dir(repo_root: Path) -> Path:
    env_values = load_env_file(repo_root / "common.env")
    yaml_config = load_yaml_config(repo_root / "config.yaml")
    raw_input = env_values.get("INPUT_PATH") or resolve_setting(
        yaml_config.get("input_path"), env_values, "input"
    )
    input_dir = Path(raw_input)
    if not input_dir.is_absolute():
        input_dir = repo_root / input_dir
    return input_dir


def resolve_skipped_pdf_names(repo_root: Path) -> set[str]:
    env_values = load_env_file(repo_root / "common.env")
    return set(split_name_list(env_values.get("SKIP_PDF_FILES", "")))


def resolve_log_dir(repo_root: Path) -> Path:
    env_values = load_env_file(repo_root / "common.env")
    yaml_config = load_yaml_config(repo_root / "config.yaml")
    raw_log_dir = env_values.get("LOG_DIR") or resolve_setting(
        yaml_config.get("log_dir"), env_values, "log"
    )
    log_dir = Path(raw_log_dir)
    if not log_dir.is_absolute():
        log_dir = repo_root / log_dir
    return log_dir


def setup_logging(repo_root: Path, script_name: str | None = None) -> Path:
    log_dir = resolve_log_dir(repo_root)
    log_dir.mkdir(parents=True, exist_ok=True)
    log_stem = script_name or Path(sys.argv[0]).stem or Path(__file__).stem
    log_stem = re.sub(r'[\\/:*?"<>|]+', "_", log_stem)
    log_path = log_dir / f"{log_stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    log_handle = log_path.open("a", encoding="utf-8")
    sys.stdout = TeeStream(sys.stdout, log_handle)  # type: ignore[assignment]
    sys.stderr = TeeStream(sys.stderr, log_handle)  # type: ignore[assignment]
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%H:%M:%S",
        stream=sys.stdout,
        force=True,
    )
    return log_path


def main() -> int:
    repo_root = Path(__file__).resolve().parent
    log_path = setup_logging(repo_root)
    parser = argparse.ArgumentParser(
        description="重命名质保作业申请单并导出 Excel 汇总"
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        help="输入目录，未指定时优先读取 common.env 中的 INPUT_PATH",
    )
    args = parser.parse_args()

    LOGGER.info("终端输出日志: %s", log_path)
    input_dir = (
        args.input_dir.resolve() if args.input_dir else resolve_input_dir(repo_root)
    )

    if not input_dir.exists():
        raise FileNotFoundError(f"输入目录不存在: {input_dir}")

    LOGGER.info("输入目录: %s", input_dir)
    skipped_pdf_names = resolve_skipped_pdf_names(repo_root)
    output_path = input_dir / SUMMARY_EXCEL_NAME
    pdf_renamed_count = rename_subject_warranty_pdfs_by_local_ai(
        input_dir, repo_root, skipped_pdf_names
    )
    records = process_documents(input_dir, output_path, skipped_pdf_names, repo_root)
    image_renamed_count = rename_matched_images_by_local_ai(input_dir, records, repo_root)
    manual_pdf_match_count = apply_manual_pdf_matches(records, input_dir, repo_root)
    LOGGER.info("开始导出 Excel: %s", output_path)
    saved_path = export_excel(records, output_path)
    LOGGER.info("Excel 导出完成: %s", saved_path)

    summary = {
        "input_dir": str(input_dir),
        "total_docs": len(records),
        "renamed_docs": sum(1 for record in records if record.处理状态 == "已重命名"),
        "renamed_pdfs": pdf_renamed_count,
        "renamed_images": image_renamed_count,
        "manual_pdf_matches": manual_pdf_match_count,
        "excel_path": str(saved_path),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
